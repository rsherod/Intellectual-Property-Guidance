# Intellectual Property Licensing Agreement Assistant
import streamlit as st
import google.generativeai as genai
from PIL import Image
import json
import os
from datetime import datetime
from io import BytesIO
import PyPDF2
import docx

# Import file format libraries (will be used in helper functions)
# Note: You may need to add these to your requirements.txt
# python-docx, reportlab, PyPDF2

# Streamlit configuration
st.set_page_config(page_title="IP Licensing Agreement Assistant", layout="wide")

# Initialize session state variables
if "messages" not in st.session_state:
    st.session_state.messages = []
if "model_name" not in st.session_state:
    st.session_state.model_name = "gemini-2.0-flash"
if "temperature" not in st.session_state:
    st.session_state.temperature = 0.3  # Lower temperature for more consistent legal language
if "debug" not in st.session_state:
    st.session_state.debug = []
if "chat_session" not in st.session_state:
    st.session_state.chat_session = None
if "uploaded_policies" not in st.session_state:
    st.session_state.uploaded_policies = {}
if "policy_analysis" not in st.session_state:
    st.session_state.policy_analysis = ""

# Helper function to extract text from PDF
def extract_text_from_pdf(uploaded_file):
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

# Helper function to extract text from Word document
def extract_text_from_docx(uploaded_file):
    try:
        doc = docx.Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {e}")
        return ""

# Helper function to load text files
def load_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        st.error(f"Error loading text file: {e}")
        return ""

# Helper function to load JSON files
def load_json_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    except Exception as e:
        st.error(f"Error loading JSON file: {e}")
        return {}

# Helper functions for different file format exports
def get_chat_text_markdown():
    """Convert the chat messages to a downloadable markdown format"""
    chat_text = "# IP Licensing Agreement Assistant Chat Log\n\n"
    
    # Add timestamp
    chat_text += f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    
    # Add uploaded policies info if applicable
    if st.session_state.uploaded_policies:
        chat_text += "## Uploaded Policy Documents:\n"
        for filename in st.session_state.uploaded_policies.keys():
            chat_text += f"- {filename}\n"
        chat_text += "\n"
    
    # Add the messages
    for msg in st.session_state.messages:
        role = "User" if msg["role"] == "user" else "Assistant"
        chat_text += f"## {role}:\n{msg['content']}\n\n"
    
    return chat_text

def get_chat_pdf():
    """Convert the chat messages to a PDF file using reportlab"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from io import BytesIO
    
    # Create a buffer for the PDF
    buffer = BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create a list to hold the elements
    elements = []
    
    # Add title
    title_style = styles["Title"]
    elements.append(Paragraph("IP Licensing Agreement Assistant Chat Log", title_style))
    elements.append(Spacer(1, 12))
    
    # Add timestamp
    normal_style = styles["Normal"]
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    elements.append(Spacer(1, 12))
    
    # Add uploaded policies info if applicable
    if st.session_state.uploaded_policies:
        elements.append(Paragraph("Uploaded Policy Documents:", styles["Heading2"]))
        for filename in st.session_state.uploaded_policies.keys():
            elements.append(Paragraph(f"• {filename}", normal_style))
        elements.append(Spacer(1, 12))
    
    # Define styles for user and assistant
    user_style = ParagraphStyle(
        "UserStyle", 
        parent=styles["Heading2"],
        textColor=colors.darkblue
    )
    
    assistant_style = ParagraphStyle(
        "AssistantStyle", 
        parent=styles["Heading2"],
        textColor=colors.darkgreen
    )
    
    # Add the messages
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            role = "User"
            heading_style = user_style
        else:
            role = "Assistant"
            heading_style = assistant_style
        
        # Add role heading
        elements.append(Paragraph(f"{role}:", heading_style))
        elements.append(Spacer(1, 6))
        
        # Add message content with proper line breaks
        text = msg["content"].replace('\n', '<br/>')
        elements.append(Paragraph(text, normal_style))
        elements.append(Spacer(1, 12))
    
    # Build the PDF
    doc.build(elements)
    
    # Get the value from the buffer
    pdf_value = buffer.getvalue()
    buffer.close()
    
    return pdf_value

def get_chat_docx():
    """Convert the chat messages to a Word document"""
    from docx import Document
    from io import BytesIO
    
    # Create document
    doc = Document()
    
    # Add title
    doc.add_heading('IP Licensing Agreement Assistant Chat Log', 0)
    
    # Add timestamp
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add uploaded policies info if applicable
    if st.session_state.uploaded_policies:
        doc.add_heading('Uploaded Policy Documents:', 2)
        for filename in st.session_state.uploaded_policies.keys():
            doc.add_paragraph(f"• {filename}")
    
    # Add the messages
    for msg in st.session_state.messages:
        role = "User" if msg["role"] == "user" else "Assistant"
        
        # Add role as heading
        doc.add_heading(f"{role}:", 2)
        
        # Add message content
        doc.add_paragraph(msg["content"])
    
    # Save to BytesIO object
    bytesio = BytesIO()
    doc.save(bytesio)
    bytesio.seek(0)
    
    # Return bytes
    return bytesio.getvalue()

# Function to save messages (removed Firestore functionality)
def save_message(user_input, bot_response):
    # Firebase functionality removed - messages are only stored in session state
    st.session_state.debug.append("Message saved to session state only")

# Load pre-existing policy examples
def load_policy_examples():
    policy_examples = {}
    
    # Load the Policy_Examples.json if it exists
    if os.path.exists('Policy_Examples.json'):
        policy_examples = load_json_file('Policy_Examples.json')
        st.session_state.debug.append("Loaded pre-existing policy examples")
    else:
        st.session_state.debug.append("Warning: Policy_Examples.json not found")
    
    return policy_examples

# Function to build the complete system prompt
def build_system_prompt():
    # Load the system instructions from the documents
    system_instructions = load_text_file('paste-2.txt')
    
    # Load pre-existing policy examples
    policy_examples = load_policy_examples()
    
    prompt = system_instructions
    
    # Add pre-loaded policy examples
    if policy_examples:
        prompt += "\n\n## Pre-Loaded Policy Examples\n\n"
        prompt += json.dumps(policy_examples, indent=2)
    
    # Add user-uploaded policies if any
    if st.session_state.uploaded_policies:
        prompt += "\n\n## User-Uploaded Policy Documents\n\n"
        for filename, content in st.session_state.uploaded_policies.items():
            prompt += f"### {filename}\n\n"
            prompt += content + "\n\n"
    
    return prompt

# Main app layout
def main():
    # Initialize Gemini (Firebase removed)
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
    except KeyError:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            st.error("⚠️ Google API Key not found!")
            st.stop()

    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

    # Header and description
    st.markdown("<h1 style='text-align: center; color: #2E86AB;'>IP Licensing Agreement Assistant</h1>", unsafe_allow_html=True)
    
    st.markdown("""
    <div style='background-color: #F0F8FF; padding: 20px; border-radius: 10px; margin-bottom: 20px;'>
    <h3>Welcome to the IP Licensing Agreement Assistant!</h3>
    <p>I'm here to help you draft intellectual property licensing agreements for researchers and faculty members. I can:</p>
    <ul>
        <li>Create detailed licensing agreement sections and clauses</li>
        <li>Analyze your institutional IP policies</li>
        <li>Handle multiple policy sources and identify conflicts</li>
        <li>Provide comprehensive coverage of IP rights and protections</li>
    </ul>
    <p><strong>To get started:</strong> Upload any relevant institutional policy documents, then describe what type of licensing agreement you need.</p>
    </div>
    """, unsafe_allow_html=True)

    # File upload section
    st.markdown("### 📄 Upload Institutional Policy Documents (Optional)")
    uploaded_files = st.file_uploader(
        "Upload your institutional IP policies (PDF, Word, or text files)",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        help="Upload any relevant institutional policies that should govern the licensing agreement"
    )

    # Process uploaded files
    if uploaded_files:
        for uploaded_file in uploaded_files:
            if uploaded_file.name not in st.session_state.uploaded_policies:
                with st.spinner(f"Processing {uploaded_file.name}..."):
                    file_content = ""
                    
                    if uploaded_file.type == "application/pdf":
                        file_content = extract_text_from_pdf(uploaded_file)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        file_content = extract_text_from_docx(uploaded_file)
                    elif uploaded_file.type == "text/plain":
                        file_content = str(uploaded_file.read(), "utf-8")
                    
                    if file_content:
                        st.session_state.uploaded_policies[uploaded_file.name] = file_content
                        st.success(f"✅ Successfully processed {uploaded_file.name}")
                        # Reset chat session to incorporate new policy
                        st.session_state.chat_session = None

    # Display uploaded policies
    if st.session_state.uploaded_policies:
        st.markdown("#### Uploaded Policy Documents:")
        for filename in st.session_state.uploaded_policies.keys():
            col1, col2 = st.columns([4, 1])
            with col1:
                st.write(f"📋 {filename}")
            with col2:
                if st.button("Remove", key=f"remove_{filename}"):
                    del st.session_state.uploaded_policies[filename]
                    st.session_state.chat_session = None  # Reset to incorporate changes
                    st.rerun()

    st.divider()

    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Download chat functionality
    if st.session_state.messages:
        st.markdown("---")
        download_col1, download_col2, download_col3, download_col4, download_col5 = st.columns([2, 2, 2, 2, 2])
        
        with download_col3:
            format_option = st.selectbox(
                "Download Format:",
                ["Markdown (.md)", "PDF (.pdf)", "Word (.docx)"],
                key="format_selection"
            )
        
        with download_col4:
            # Determine file extension and data based on format selection
            if format_option == "PDF (.pdf)":
                file_ext = ".pdf"
                file_data = get_chat_pdf()
                mime_type = "application/pdf"
            elif format_option == "Word (.docx)":
                file_ext = ".docx"
                file_data = get_chat_docx()
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:  # Default to Markdown
                file_ext = ".md"
                file_data = get_chat_text_markdown()
                mime_type = "text/markdown"
            
            base_filename = f"IP-Licensing-Chat-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            
            st.download_button(
                label="💾 Download Chat",
                data=file_data,
                file_name=f"{base_filename}{file_ext}",
                mime=mime_type,
                help="Save this conversation to your device"
            )

    # Chat input
    user_input = st.chat_input("Describe the licensing agreement you need help with...")

    if user_input:
        # Add user message to chat history
        current_message = {"role": "user", "content": user_input}
        st.session_state.messages.append(current_message)

        with st.chat_message("user"):
            st.markdown(current_message["content"])

        # Generate and display assistant response
        with st.chat_message("assistant"):
            message_placeholder = st.empty()

            # Initialize chat session if needed
            if st.session_state.chat_session is None:
                generation_config = {
                    "temperature": st.session_state.temperature,
                    "top_p": 0.95,
                    "top_k": 40,
                    "max_output_tokens": 8192,
                }
                model = genai.GenerativeModel(
                    model_name=st.session_state.model_name,
                    generation_config=generation_config,
                )
                
                # Build complete system prompt
                complete_system_prompt = build_system_prompt()
                
                # Initialize chat with system prompt
                initial_messages = [
                    {"role": "user", "parts": [f"System: {complete_system_prompt}"]},
                    {"role": "model", "parts": ["Understood. I am now configured as an IP Licensing Agreement Assistant and will help you draft licensing agreements based on institutional policies and your specific requirements."]},
                ]
                
                st.session_state.chat_session = model.start_chat(history=initial_messages)

            # Generate response with error handling
            try:
                response = st.session_state.chat_session.send_message(current_message["content"])
                full_response = response.text
                message_placeholder.markdown(full_response)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                st.session_state.debug.append("Assistant response generated")
                save_message(current_message["content"], full_response)
            except Exception as e:
                st.error(f"An error occurred while generating the response: {e}")
                st.session_state.debug.append(f"Error: {e}")

        st.rerun()

    # Footer with legal disclaimer
    st.markdown("---")
    st.markdown("""
    <div style='background-color: #FFF3CD; padding: 15px; border-radius: 5px; border-left: 4px solid #FFC107;'>
    <p><strong>⚠️ IMPORTANT LEGAL DISCLAIMER:</strong></p>
    <p style='font-size: 0.9em;'>This assistant provides draft language for informational purposes only and does not constitute legal advice. All licensing agreements must be reviewed and approved by appropriate institutional authorities, including university technology transfer offices, legal counsel, and institutional IP officials. Please ensure all drafts undergo proper institutional review and approval before execution.</p>
    </div>
    """, unsafe_allow_html=True)

    # Debug information (only show if there are debug messages)
    if st.session_state.debug and st.checkbox("Show Debug Information"):
        st.markdown("### Debug Information")
        for debug_msg in st.session_state.debug:
            st.text(debug_msg)

if __name__ == "__main__":
    main()
